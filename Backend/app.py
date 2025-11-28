import os
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, AsyncGenerator
import asyncio
import json
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.output_parsers import StrOutputParser
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime as dt
from pathlib import Path
import re
from dotenv import load_dotenv
from db_connect import get_database, close_database
from models import SRSDocument, SRSRepository

# Load environment variables
load_dotenv()

app = FastAPI(title="SRS Generation API")

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "https://blueprint-ai-platinum.vercel.app"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Base directories for document storage
BASE_STORAGE_DIR = Path(__file__).parent / "storage"
BASE_STORAGE_DIR.mkdir(exist_ok=True)

# MongoDB instance
db = None
srs_repo = None

@app.on_event("startup")
async def startup_db_client():
    """Initialize MongoDB connection on startup"""
    global db, srs_repo
    try:
        db = get_database()
        srs_repo = SRSRepository(db)
        print("[+] Database initialized")
    except Exception as e:
        print(f"[-] Failed to initialize database: {e}")

@app.on_event("shutdown")
async def shutdown_db_client():
    """Close MongoDB connection on shutdown"""
    close_database()
    print("[+] Application shutdown complete")

# Initialize OpenAI with LangChain
def get_llm():
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY environment variable is not set")
    return ChatOpenAI(
        model="gpt-4o-mini", 
        temperature=0.8,
        max_tokens=16384,
        api_key=api_key
    )

# Pydantic Models
class SRSGenerationRequest(BaseModel):
    main: str
    selectedPurpose: Optional[str] = ""
    selectedTarget: Optional[str] = ""
    selectedKeys: Optional[str] = ""
    selectedPlatforms: Optional[str] = ""
    selectedIntegrations: Optional[str] = ""
    selectedPerformance: Optional[str] = ""
    selectedSecurity: Optional[str] = ""
    selectedStorage: Optional[str] = ""
    selectedEnvironment: Optional[str] = ""
    selectedLanguage: Optional[str] = ""
    userId: Optional[str] = None  # MongoDB user ID
    username: Optional[str] = None  # Username for file storage

class PDFGenerationRequest(BaseModel):
    username: str
    text: str
    title: str

# Helper Functions
def get_user_directories(username: str) -> tuple[Path, Path]:
    """Get or create user-specific directories for PDFs and Word documents"""
    user_dir = BASE_STORAGE_DIR / username
    pdfs_dir = user_dir / "pdfs"
    docs_dir = user_dir / "docs"
    
    # Create directories if they don't exist
    pdfs_dir.mkdir(parents=True, exist_ok=True)
    docs_dir.mkdir(parents=True, exist_ok=True)
    
    return pdfs_dir, docs_dir

def sanitize_filename(filename: str) -> str:
    """Sanitize filename to prevent directory traversal and other security issues"""
    # Remove any directory separators and parent directory references
    filename = os.path.basename(filename)
    # Remove any non-alphanumeric characters except dots, hyphens, and underscores
    filename = re.sub(r'[^\w\-\.]', '_', filename)
    return filename

def create_pdf(title: str, text: str, username: str) -> tuple[str, str]:
    """Generate PDF document using reportlab"""
    try:
        # Get user-specific directories
        pdfs_dir, _ = get_user_directories(username)
        
        # Create filename with timestamp
        timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{username}_{timestamp}.pdf"
        filepath = pdfs_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1a1a1a',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            textColor='#333333',
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            fontName='Helvetica'
        )
        
        # Heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#1a1a1a',
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        # Add title
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Add metadata
        # metadata = f"Generated by: {username}<br/>Date: {dt.datetime.now().strftime('%B %d, %Y')}"
        # elements.append(Paragraph(metadata, styles['Normal']))
        # elements.append(Spacer(1, 0.3*inch))
        
        # Process text content
        # Split by newlines and process each line
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
            
            # Check if line is a heading (starts with numbers like 1., 1.1, etc.)
            if re.match(r'^\d+\.', line):
                elements.append(Paragraph(line, heading_style))
            else:
                # Escape special characters for ReportLab
                line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                elements.append(Paragraph(line, body_style))
        
        # Build PDF
        doc.build(elements)
        
        # Return both filename and relative path from username
        relative_path = f"{username}/pdfs/{filename}"
        return filename, relative_path
    except Exception as e:
        print(f"Error creating PDF: {e}")
        raise

def create_word(title: str, text: str, username: str) -> tuple[str, str]:
    """Generate Word document using python-docx"""
    try:
        # Get user-specific directories
        _, docs_dir = get_user_directories(username)
        
        # Create filename with timestamp
        timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{username}_{timestamp}.docx"
        filepath = docs_dir / filename
        
        # Create Document
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        # metadata_para = doc.add_paragraph()
        # metadata_para.add_run(f"Generated by: {username}\n").bold = True
        # metadata_para.add_run(f"Date: {dt.datetime.now().strftime('%B %d, %Y')}")
        # metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Process text content
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()  # Empty line
                continue
            
            # Check if line is a heading
            if re.match(r'^\d+\.', line):
                heading = doc.add_heading(line, level=1)
                heading_format = heading.runs[0].font
                heading_format.size = Pt(14)
            else:
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_format = para.runs[0].font if para.runs else None
                if para_format:
                    para_format.size = Pt(11)
        
        # Save document
        doc.save(str(filepath))
        
        # Return both filename and relative path from username
        relative_path = f"{username}/docs/{filename}"
        return filename, relative_path
    except Exception as e:
        print(f"Error creating Word document: {e}")
        raise

# API Routes
@app.get("/")
async def root():
    return {"message": "[+] Server up and running..."}

@app.post("/generate-srs")
async def generate_srs(request: SRSGenerationRequest):
    """Generate SRS document using LangChain and OpenAI"""
    try:
        # Initialize LangChain
        llm = get_llm()
        output_parser = StrOutputParser()
        
        # Create the prompt
        prompt = f"""
        Generate a comprehensive Software Requirements Specification (SRS) document with the following details:
        
        Main Idea: {request.main}
        Primary Purpose: {request.selectedPurpose}
        Target Users: {request.selectedTarget}
        Key Features: {request.selectedKeys}
        Compatible Platforms: {request.selectedPlatforms}
        Integration Requirements: {request.selectedIntegrations}
        Performance Requirements: {request.selectedPerformance}
        Security Requirements: {request.selectedSecurity}
        Data Storage Capacity: {request.selectedStorage}
        Operating Environment: {request.selectedEnvironment}
        Language and Localization: {request.selectedLanguage}
        
        Please provide a suitable title for the software based on the main idea and write it like "Title: [title generated]" at the top of the text.
        
        Format the document with proper sections including:
        1. Introduction
        2. Overall Description
        3. Specific Requirements
        4. System Features
        5. External Interface Requirements
        6. Non-Functional Requirements
        
        Use newline characters for formatting. Do not use markdown hash symbols (#) for headings.
        Write in a professional, technical style appropriate for an SRS document.
        """
        
        # Create messages
        messages = [
            SystemMessage(content="You are an expert technical writer specializing in Software Requirements Specification (SRS) documents. You generate comprehensive, well-structured SRS documents following IEEE 830 standards."),
            HumanMessage(content=prompt)
        ]
        
        # Generate content
        response = llm.invoke(messages)
        generated_text = output_parser.invoke(response)
        
        # Extract title
        title_match = re.search(r'Title:\s*(.*)', generated_text)
        title = title_match.group(1).strip() if title_match else "SRS DOCUMENT"
        # Remove asterisks and other markdown formatting
        title = re.sub(r'[\*\#\_]', '', title)
        
        # Remove the "Title:" line from the text
        modified_text = re.sub(r'Title:\s*.*\n?', '', generated_text, count=1)
        
        return JSONResponse({
            "success": True,
            "title": title,
            "text": modified_text,
            "message": "SRS generated successfully"
        })
        
    except Exception as e:
        print(f"Error generating SRS: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error generating SRS: {str(e)}"
        )

@app.post("/generate-pdf")
async def generate_pdf(request: PDFGenerationRequest):
    """Generate PDF and Word documents from text"""
    try:
        if not request.username:
            raise HTTPException(
                status_code=400,
                detail="Username is required"
            )
        
        # Generate PDF
        pdf_filename, pdf_path = create_pdf(request.title, request.text, request.username)
        
        # Generate Word document
        word_filename, word_path = create_word(request.title, request.text, request.username)
        
        return JSONResponse({
            "success": True,
            "pdfName": pdf_filename,
            "wordName": word_filename,
            "pdfPath": pdf_path,
            "wordPath": word_path,
            "message": "Documents generated successfully"
        })
        
    except Exception as e:
        print(f"Error generating documents: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error generating documents: {str(e)}"
        )

@app.get("/download-pdf/{username}/{filename}")
async def download_pdf(username: str, filename: str):
    """Download PDF file from user's directory"""
    try:
        # Sanitize inputs
        safe_username = sanitize_filename(username)
        safe_filename = sanitize_filename(filename)
        
        pdfs_dir, _ = get_user_directories(safe_username)
        filepath = pdfs_dir / safe_filename
        
        if not filepath.exists():
            raise HTTPException(
                status_code=404,
                detail="PDF file not found"
            )
        
        return FileResponse(
            path=str(filepath),
            filename=safe_filename,
            media_type='application/pdf'
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error downloading PDF: {e}")
        raise HTTPException(
            status_code=500,
            detail="An error occurred while downloading the PDF file"
        )

@app.get("/download-word/{username}/{filename}")
async def download_word(username: str, filename: str):
    """Download Word document from user's directory"""
    try:
        # Sanitize inputs
        safe_username = sanitize_filename(username)
        safe_filename = sanitize_filename(filename)
        
        _, docs_dir = get_user_directories(safe_username)
        filepath = docs_dir / safe_filename
        
        if not filepath.exists():
            raise HTTPException(
                status_code=404,
                detail="Word document not found"
            )
        
        return FileResponse(
            path=str(filepath),
            filename=safe_filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error downloading Word document: {e}")
        raise HTTPException(
            status_code=500,
            detail="An error occurred while downloading the Word document"
        )

async def srs_generation_stream(request: SRSGenerationRequest) -> AsyncGenerator[str, None]:
    """
    Generator function for Server-Sent Events
    Streams progress updates during SRS generation
    """
    srs_id = None
    try:
        # Send initial status
        yield f"data: {json.dumps({'status': 'initiated', 'message': 'Starting SRS generation...'})}\n\n"
        await asyncio.sleep(0.1)  # Small delay to ensure message is sent
        # Determine username for file storage
        username = request.username
        user_id = request.userId
        
        # Create description array similar to Next.js implementation
        description_array = [
            request.main,
            request.selectedPurpose,
            request.selectedTarget,
            request.selectedKeys,
            request.selectedPlatforms,
            request.selectedIntegrations,
            request.selectedPerformance,
            request.selectedSecurity,
            request.selectedStorage,
            request.selectedEnvironment,
            request.selectedLanguage,
        ]
        
        # Save initial SRS document to database with "Processing" status
        if srs_repo:
            try:
                initial_srs = SRSDocument(
                    owner=user_id,
                    name="Generating...",  # Will be updated with actual title
                    description=json.dumps(description_array),
                    status="Processing",
                    pdf_url="",
                    word_url=""
                )
                srs_id = srs_repo.create(initial_srs)
                yield f"data: {json.dumps({'status': 'processing', 'message': 'SRS record created in database...'})}\n\n"
                await asyncio.sleep(0.1)  # Ensure message is flushed
            except Exception as db_error:
                print(f"Database error: {db_error}")
                # Continue even if database save fails
                pass
        
        # Initialize LangChain
        yield f"data: {json.dumps({'status': 'processing', 'message': 'Initializing AI model...'})}\n\n"
        await asyncio.sleep(0.1)  # Ensure message is flushed
        
        llm = get_llm()
        output_parser = StrOutputParser()
        
        # Create the prompt
        prompt = f"""
        Generate a comprehensive Software Requirements Specification (SRS) document with the following details:
        
        Main Idea: {request.main}
        Primary Purpose: {request.selectedPurpose}
        Target Users: {request.selectedTarget}
        Key Features: {request.selectedKeys}
        Compatible Platforms: {request.selectedPlatforms}
        Integration Requirements: {request.selectedIntegrations}
        Performance Requirements: {request.selectedPerformance}
        Security Requirements: {request.selectedSecurity}
        Data Storage Capacity: {request.selectedStorage}
        Operating Environment: {request.selectedEnvironment}
        Language and Localization: {request.selectedLanguage}
        
        Please provide a suitable title for the software based on the main idea and write it like "Title: [title generated]" at the top of the text.
        
        Format the document with proper sections including:
        1. Introduction
        2. Overall Description
        3. Specific Requirements
        4. System Features
        5. External Interface Requirements
        6. Non-Functional Requirements
        
        Use newline characters for formatting. Do not use markdown hash symbols (#) for headings.
        Write in a professional, technical style appropriate for an SRS document.
        """
        
        # Generate content
        yield f"data: {json.dumps({'status': 'processing', 'message': 'Generating SRS content with AI...'})}\n\n"
        await asyncio.sleep(0.1)  # Ensure message is flushed before long AI call
        
        messages = [
            SystemMessage(content="You are an expert technical writer specializing in Software Requirements Specification (SRS) documents. You generate comprehensive, well-structured SRS documents following IEEE 830 standards."),
            HumanMessage(content=prompt)
        ]
        
        response = llm.invoke(messages)
        generated_text = output_parser.invoke(response)
        
        # Extract title
        yield f"data: {json.dumps({'status': 'processing', 'message': 'Processing generated content...'})}\n\n"
        await asyncio.sleep(0.1)  # Ensure message is flushed
        
        title_match = re.search(r'Title:\s*(.*)', generated_text)
        title = title_match.group(1).strip() if title_match else "SRS DOCUMENT"
        title = re.sub(r'[\*\#\_]', '', title)
        
        # Remove the "Title:" line from the text
        modified_text = re.sub(r'Title:\s*.*\n?', '', generated_text, count=1)
        
        yield f"data: {json.dumps({'status': 'processing', 'message': f'SRS generated: {title}', 'title': title})}\n\n"
        await asyncio.sleep(0.1)  # Ensure message is flushed
        
        # Generate PDF
        yield f"data: {json.dumps({'status': 'processing', 'message': 'Creating PDF document...'})}\n\n"
        await asyncio.sleep(0.1)  # Ensure message is flushed
        
        pdf_filename, pdf_path = create_pdf(title, modified_text, username)
        
        # Generate Word document
        yield f"data: {json.dumps({'status': 'processing', 'message': 'Creating Word document...'})}\n\n"
        await asyncio.sleep(0.1)  # Ensure message is flushed
        
        word_filename, word_path = create_word(title, modified_text, username)
        
        # Update database with completion status and file URLs
        if srs_repo and srs_id:
            try:
                srs_repo.update(srs_id, {
                    "name": title,
                    "status": "Completed",
                    "pdf_url": pdf_filename,  # Store just filename like Next.js
                    "word_url": word_filename
                })
                yield f"data: {json.dumps({'status': 'processing', 'message': 'Database updated with generated files...'})}\n\n"
                await asyncio.sleep(0.1)  # Ensure message is flushed
            except Exception as db_error:
                print(f"Database update error: {db_error}")
                # Continue even if database update fails
                pass
        
        # Send completion event with file information
        completion_data = {
            'status': 'completed',
            'message': 'SRS generation completed successfully!',
            'title': title,
            'pdfName': pdf_filename,
            'wordName': word_filename,
            'pdfPath': pdf_path,
            'wordPath': word_path,
            'text': modified_text,
            'srsId': srs_id  # Include database ID
        }
        yield f"data: {json.dumps(completion_data)}\n\n"
        
        # Send close event
        yield "event: close\ndata: {}\n\n"
        
    except Exception as e:
        # Update database with failed status if SRS was created
        if srs_repo and srs_id:
            try:
                srs_repo.update(srs_id, {
                    "status": "Failed",
                    "pdf_url": "No PDF",
                    "word_url": "No Docx"
                })
            except Exception as db_error:
                print(f"Database error during failure handling: {db_error}")
        
        error_data = {
            'status': 'error',
            'message': f'Error during SRS generation: {str(e)}'
        }
        yield f"data: {json.dumps(error_data)}\n\n"
        yield "event: close\ndata: {}\n\n"

@app.post("/generate-srs-stream")
async def generate_srs_stream(request: SRSGenerationRequest):
    """
    Generate SRS document with real-time progress updates via Server-Sent Events (SSE)
    This endpoint combines SRS generation and document creation in one flow
    """
    return StreamingResponse(
        srs_generation_stream(request),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"  # Disable buffering for nginx
        }
    )

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run("app:app", host="0.0.0.0", port=port, reload=True)
